import os
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
import PyPDF2
from docx import Document
from app.core.logging import logger
from app.core.config import settings


class DocumentParser:
    """Document parser for PDF, DOCX, and TXT files"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
        self.max_file_size = settings.max_file_size * 1024 * 1024  # Convert MB to bytes
    
    def validate_file(self, file_path: str) -> tuple[bool, str]:
        """Validate file format and size"""
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                return False, "File does not exist"
            
            # Check file extension
            if path.suffix.lower() not in self.supported_formats:
                return False, f"Unsupported format. Supported formats: {', '.join(self.supported_formats)}"
            
            # Check file size
            file_size = path.stat().st_size
            if file_size > self.max_file_size:
                return False, f"File too large. Maximum size: {settings.max_file_size}MB"
            
            return True, "File is valid"
            
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return False, f"Error validating file: {str(e)}"
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Extracted {len(text)} characters from TXT: {file_path}")
            return text.strip()
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                logger.info(f"Extracted {len(text)} characters from TXT (latin-1): {file_path}")
                return text.strip()
                
            except Exception as e:
                logger.error(f"TXT extraction error (encoding): {e}")
                raise ValueError(f"Error reading TXT file: {str(e)}")
                
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise ValueError(f"Error extracting text from TXT: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from supported document formats"""
        # Validate file first
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(message)
        
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-\.,;:!?@#$%&*()_+=\[\]{}|\\/"\'<>]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get file information"""
        try:
            path = Path(file_path)
            stat = path.stat()
            
            return {
                'filename': path.name,
                'extension': path.suffix.lower(),
                'size_bytes': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'modified_time': stat.st_mtime,
                'is_valid': self.validate_file(file_path)[0]
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {
                'filename': Path(file_path).name,
                'error': str(e)
            }


# Global document parser instance
document_parser = DocumentParser()
