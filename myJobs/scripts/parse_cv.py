import sys
import os
import subprocess

def install_and_import(package_name, import_name=None):
    if import_name is None:
        import_name = package_name
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package_name}...", file=sys.stderr)
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
        __import__(import_name)

# Ensure packages are installed
try:
    install_and_import("pypdf")
    install_and_import("python-docx", "docx")
except Exception as e:
    print(f"Failed to install dependencies: {e}", file=sys.stderr)

import pypdf
import docx

def parse_pdf(file_path):
    reader = pypdf.PdfReader(file_path)
    text_content = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_content.append(text)
    return "\n".join(text_content)

def parse_docx(file_path):
    doc = docx.Document(file_path)
    text_content = []
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_txt = cell.text.strip()
                if cell_txt and cell_txt not in row_text: # avoid duplication in merged cells
                    row_text.append(cell_txt)
            if row_text:
                text_content.append(" | ".join(row_text))
                
    return "\n".join(text_content)

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 parse_cv.py <path_to_cv_file>", file=sys.stderr)
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
        
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            text = parse_pdf(file_path)
        elif ext == ".docx":
            text = parse_docx(file_path)
        elif ext in [".txt", ".md", ".markdown"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            print(f"Unsupported file format: {ext}", file=sys.stderr)
            sys.exit(1)
            
        print(text)
    except Exception as e:
        print(f"Error parsing file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
